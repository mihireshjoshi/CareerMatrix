import os
import random
import string
import docx
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx2pdf import convert

# Function to prompt user for input
def get_user_input(field_name):
    return input(f"Enter your {field_name}: ")

# Function to generate the resume
def generate_resume():
    # Prompt user for personal information
    name = get_user_input("full name")
    email = get_user_input("email")
    phone = get_user_input("phone number")
    address = get_user_input("address")

    # Prompt user for career objective
    career_objective = get_user_input("career objective")

    # Prompt user for qualifications
    qualifications = []
    num_qualifications = int(input("Enter the number of qualifications: "))
    for i in range(num_qualifications):
        qualification = get_user_input(f"qualification {i+1}")
        institution = get_user_input(f"institution for qualification {i+1}")
        qualifications.append({'qualification': qualification, 'institution': institution})

    # Prompt user for co-curricular skills
    co_curricular_skills = []
    num_skills = int(input("Enter the number of co-curricular skills: "))
    for i in range(num_skills):
        skill = get_user_input(f"co-curricular skill {i+1}")
        co_curricular_skills.append(skill)

    # Prompt user for work experience
    work_experience = []
    num_experiences = int(input("Enter the number of work experiences: "))
    for i in range(num_experiences):
        title = get_user_input(f"title of work experience {i+1}")
        company = get_user_input(f"company of work experience {i+1}")
        duration = get_user_input(f"duration of work experience {i+1}")
        description = get_user_input(f"description of work experience {i+1}")
        work_experience.append({'title': title, 'company': company, 'duration': duration, 'description': description})

    # Prompt user for projects
    projects = []
    num_projects = int(input("Enter the number of projects: "))
    for i in range(num_projects):
        title = get_user_input(f"title of project {i+1}")
        description = get_user_input(f"description of project {i+1}")
        projects.append({'title': title, 'description': description})

    # Prompt user for internships
    internships = []
    num_internships = int(input("Enter the number of internships: "))
    for i in range(num_internships):
        title = get_user_input(f"title of internship {i+1}")
        company = get_user_input(f"company of internship {i+1}")
        duration = get_user_input(f"duration of internship {i+1}")
        description = get_user_input(f"description of internship {i+1}")
        internships.append({'title': title, 'company': company, 'duration': duration, 'description': description})

    # Prompt user for awards or honors
    awards = []
    num_awards = int(input("Enter the number of awards or honors: "))
    for i in range(num_awards):
        title = get_user_input(f"title of award or honor {i+1}")
        issuer = get_user_input(f"issuer of award or honor {i+1}")
        date = get_user_input(f"date of award or honor {i+1}")
        awards.append({'title': title, 'issuer': issuer, 'date': date})

    # Prompt user for specific skills
    specific_skills = get_user_input("specific skills")

    # Create a new Word document
    doc = docx.Document()

    # Add a title
    title = doc.add_heading('RESUME', level=1)
    title.bold = True
    title.style.font.size = Pt(24)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add personal information
    doc.add_heading('Personal Information', level=2).bold = True
    doc.add_paragraph(f'Name: {name}')
    doc.add_paragraph(f'Email: {email}')
    doc.add_paragraph(f'Phone: {phone}')
    doc.add_paragraph(f'Address: {address}')

    # Add career objective
    doc.add_heading('Career Objective', level=2).bold = True
    p = doc.add_paragraph(career_objective)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p.style.font.size = Pt(14)

    # Add qualifications
    doc.add_heading('Qualifications', level=2).bold = True
    for qual in qualifications:
        doc.add_paragraph(f"Qualification: {qual['qualification']}")
        doc.add_paragraph(f"Institution: {qual['institution']}")

    # Add co-curricular skills
    doc.add_heading('Co-curricular Skills', level=2).bold = True
    for skill in co_curricular_skills:
        doc.add_paragraph(skill)

    # Add work experience
    doc.add_heading('Work Experience', level=2).bold = True
    for exp in work_experience:
        p = doc.add_paragraph()
        p.add_run(exp['title']).bold = True
        p.add_run(f" ({exp['duration']})")
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        doc.add_paragraph(f"Company: {exp['company']}")
        doc.add_paragraph(exp['description'])

    # Add projects
    doc.add_heading('Projects', level=2).bold = True
    for project in projects:
        p = doc.add_paragraph()
        p.add_run(project['title']).bold = True
        doc.add_paragraph(project['description'])

    # Add internships
    doc.add_heading('Internships', level=2).bold = True
    for internship in internships:
        p = doc.add_paragraph()
        p.add_run(internship['title']).bold = True
        p.add_run(f" ({internship['duration']})")
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        doc.add_paragraph(f"Company: {internship['company']}")
        doc.add_paragraph(internship['description'])

    # Add awards or honors
    doc.add_heading('Awards or Honors', level=2).bold = True
    for award in awards:
        p = doc.add_paragraph()
        p.add_run(award['title']).bold = True
        doc.add_paragraph(f"Issuer: {award['issuer']}")
        doc.add_paragraph(f"Date: {award['date']}")

    # Add specific skills
    doc.add_heading('Specific Skills', level=2).bold = True
    p = doc.add_paragraph(specific_skills)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p.style.font.size = Pt(14)

    # Save the document
    output_path = 'resume.docx'
    doc.save(output_path)

    # Convert the document to PDF
    pdf_path = output_path.replace('.docx', '.pdf')
    convert(output_path, pdf_path)

    # Remove the intermediate Word document
    os.remove(output_path)

    return pdf_path

if __name__ == "__main__":
    # Generate the resume
    pdf_path = generate_resume()
    print(f"Resume generated: {pdf_path}")
